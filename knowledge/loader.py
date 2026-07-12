import os
from dataclasses import dataclass
from typing import Iterable, List, Optional

from langchain_core.documents import Document

from AIRAGAgent.knowledge.cleaner import clean_text
from AIRAGAgent.knowledge.cleaner import infer_document_type
from AIRAGAgent.knowledge.cleaner import infer_policy_domain
from AIRAGAgent.utils.file_handler import get_file_md5_hex, pdf_loader, txt_loader


@dataclass(frozen=True)
class SourceFile:
    path: str
    md5: Optional[str] = None


def iter_source_files(
    data_path: str,
    allowed_suffixes: Iterable[str],
    calculate_md5: bool = False,
) -> List[SourceFile]:
    suffixes = tuple(s.lower().lstrip(".") for s in allowed_suffixes)
    files: List[SourceFile] = []
    if not os.path.isdir(data_path):
        return files

    for root, _, names in os.walk(data_path):
        for name in sorted(names):
            if name.startswith("~$"):
                continue
            path = os.path.join(root, name)
            if not os.path.isfile(path):
                continue
            ext = os.path.splitext(name)[1].lower().lstrip(".")
            if ext not in suffixes:
                continue
            md5 = get_file_md5_hex(path) if calculate_md5 else None
            files.append(SourceFile(path=path, md5=md5))
    return files


def load_documents(source_file: SourceFile) -> List[Document]:
    if source_file.path.lower().endswith(".txt"):
        docs = txt_loader(source_file.path)
    elif source_file.path.lower().endswith(".pdf"):
        docs = pdf_loader(source_file.path)
    elif source_file.path.lower().endswith(".docx"):
        docs = docx_loader(source_file.path)
    elif source_file.path.lower().endswith(".doc"):
        docs = doc_loader(source_file.path)
    else:
        docs = []

    for doc in docs:
        doc.page_content = clean_text(doc.page_content)
        document_type = infer_document_type(source_file.path, doc.page_content)
        policy_domain = infer_policy_domain(source_file.path, doc.page_content)
        doc.metadata = {
            **doc.metadata,
            "source": source_file.path,
            "file_md5": source_file.md5 or "",
            "file_name": os.path.basename(source_file.path),
            "document_type": document_type,
            "policy_domain": policy_domain,
        }
    return [doc for doc in docs if doc.page_content]


def docx_loader(filepath: str) -> List[Document]:
    from docx import Document as DocxDocument

    docx = DocxDocument(filepath)
    parts = []
    for paragraph in docx.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in docx.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    content = "\n".join(parts).strip()
    if not content:
        return []
    return [Document(page_content=content, metadata={"source": filepath})]


def doc_loader(filepath: str) -> List[Document]:
    """Load legacy .doc files on Windows when Microsoft Word/pywin32 is available."""
    try:
        import pythoncom
        import win32com.client
    except ImportError as exc:
        raise RuntimeError("读取 .doc 文件需要安装 pywin32，并且本机需要可用的 Microsoft Word") from exc

    pythoncom.CoInitialize()
    word = None
    document = None
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        document = word.Documents.Open(os.path.abspath(filepath), ReadOnly=True)
        content = document.Content.Text.strip()
    finally:
        if document is not None:
            document.Close(False)
        if word is not None:
            word.Quit()
        pythoncom.CoUninitialize()

    if not content:
        return []
    return [Document(page_content=content, metadata={"source": filepath})]
